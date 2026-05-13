"""
将 LLM 生成的会议报告文本转换为格式化的 Word 文档。
识别标题层级标记（#, ##, ###, ####）映射到 Word 标题样式。
"""

import argparse
import os
import re
import sys

# re 用于 detect_heading_level

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SENSITIVE_MARKER = '【以下为原文】'


def detect_heading_level(line: str):  # -> tuple[int, str] | None:
    """检测是否为标题行，返回 (层级, 去除标记后的文本)。"""
    match = re.match(r'^(#{1,4})\s+(.+)', line)
    if match:
        level = len(match.group(1))
        text = match.group(2).strip()
        return level, text
    return None


def set_heading_style(document, paragraph, level: int):
    """设置段落为指定层级的标题样式。"""
    paragraph.style = document.styles[f'Heading {level}']

    # 调整字体大小
    font_sizes = {1: Pt(18), 2: Pt(15), 3: Pt(13), 4: Pt(12)}
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    run.font.size = font_sizes.get(level, Pt(12))
    run.font.bold = level <= 2

    # 一级标题居中
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate_docx(report_path: str, output_path: str) -> None:
    """读取报告文本，生成 Word 文档。"""
    if not os.path.exists(report_path):
        print(f"错误：文件不存在 - {report_path}", file=sys.stderr)
        sys.exit(1)

    with open(report_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    lines = content.strip().split('\n')
    in_sensitive_block = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            in_sensitive_block = False
            continue

        # 标题: 遇到标题自动结束敏感内容块
        heading = detect_heading_level(stripped)
        if heading:
            in_sensitive_block = False
            level, text = heading
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_heading_style(doc, p, level)
            continue

        # 敏感内容标记行
        if SENSITIVE_MARKER in stripped:
            in_sensitive_block = True
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.size = Pt(11)
            run.font.bold = True
            continue

        # 正文段落
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        run = p.add_run(stripped)
        run.font.size = Pt(11)

        if in_sensitive_block:
            run.font.underline = True
            run.font.bold = True

    doc.save(output_path)
    print(f"Word 文档已生成: {output_path}", file=sys.stderr)


def main():
    parser = argparse.ArgumentParser(description="将会议报告文本转换为 Word 文档")
    parser.add_argument("report_path", help="报告文本文件路径")
    parser.add_argument("output_path", nargs='?', default=None, help="输出 .docx 文件路径（默认与输入同目录、同名不同扩展名）")
    args = parser.parse_args()

    output_path = args.output_path
    if not output_path:
        base = os.path.splitext(args.report_path)[0]
        output_path = f"{base}.docx"

    generate_docx(args.report_path, output_path)


if __name__ == "__main__":
    main()
