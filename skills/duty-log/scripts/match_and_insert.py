"""Time slot matching and Word document insertion for duty log images."""
import json
from datetime import datetime, time
from pathlib import Path

# === Time Slot Matching ===

def load_timeslots():
    cfg_path = Path(__file__).parent / 'timeslots.json'
    with open(cfg_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return data['timeslots']

def parse_time(timestr):
    """Parse 'HH:MM' string to datetime.time."""
    timestr = timestr.strip()
    parts = timestr.split(':')
    return time(int(parts[0]), int(parts[1]))

def time_in_range(t, start_str, end_str):
    """Check if time t falls in [start_str, end_str] range."""
    s = parse_time(start_str)
    e = parse_time(end_str)
    if s <= e:
        return s <= t <= e
    else:
        # Overnight range
        return t >= s or t <= e

def time_in_tolerance(t, point_str, tolerance_minutes):
    """Check if time t is within tolerance_minutes of a point time."""
    p = parse_time(point_str)
    p_mins = p.hour * 60 + p.minute
    t_mins = t.hour * 60 + t.minute
    return abs(t_mins - p_mins) <= tolerance_minutes

def match_timeslot(hhmm):
    """Match an HH:MM time string to the best timeslot.

    Returns dict with section, label, content_hint or None if no match.

    Matching rules:
    1. Range slots: time must fall within [start, end].
    2. Tolerance slots: time must be within tolerance_minutes of start (bidirectional).
    3. When both a tolerance and range slot match:
       - If t >= tolerance_slot.start, tolerance wins (more specific event has started).
       - Otherwise range wins (time is still in the lead-up to the tolerance point).
    4. Within the same type, sort by earliest start time.
    5. If no active slot, look ahead up to 60 minutes for the next upcoming slot.
    """
    t = parse_time(hhmm)
    timeslots = load_timeslots()

    # Phase 1: find active slots (range or tolerance)
    candidates_range = []
    candidates_tolerance = []

    for slot in timeslots:
        if slot.get('end') is not None:
            if time_in_range(t, slot['start'], slot['end']):
                candidates_range.append(slot)
        else:
            tol = slot.get('tolerance_minutes', 10)
            if time_in_tolerance(t, slot['start'], tol):
                candidates_tolerance.append(slot)

    # Tiebreaker among active slots
    if candidates_tolerance or candidates_range:
        # Check if any tolerance slot has t >= its start time
        active_tolerance = [s for s in candidates_tolerance if t >= parse_time(s['start'])]

        if active_tolerance:
            # Tolerance slots that have "started" take priority over range slots
            active_tolerance.sort(key=lambda s: s['start'])
            return active_tolerance[0]

        if candidates_range:
            # Sort: shorter duration first (more specific), then earlier start
            def slot_duration(slot):
                if slot['end'] is None:
                    return 9999  # point slots go last
                s = parse_time(slot['start'])
                e = parse_time(slot['end'])
                return (e.hour * 60 + e.minute) - (s.hour * 60 + s.minute)

            candidates_range.sort(key=lambda s: (slot_duration(s), s['start']))
            return candidates_range[0]

        # Only unmatched tolerance (t < start for all), still return best
        candidates_tolerance.sort(key=lambda s: s['start'])
        return candidates_tolerance[0]

    # Phase 2: no active slot - look ahead for next upcoming slot within 60 minutes
    t_mins = t.hour * 60 + t.minute
    LOOKAHEAD_MINUTES = 60

    upcoming = []
    for slot in timeslots:
        slot_start = parse_time(slot['start'])
        slot_start_mins = slot_start.hour * 60 + slot_start.minute
        if slot_start_mins > t_mins:
            diff = slot_start_mins - t_mins
            upcoming.append((diff, slot))

    if upcoming:
        upcoming.sort(key=lambda x: x[0])
        closest_diff, closest_slot = upcoming[0]
        if closest_diff <= LOOKAHEAD_MINUTES:
            return closest_slot

    return None

# === Word Document Operations ===

def build_doc_name(campus, date_str):
    """Build document filename from campus and date.

    Args:
        campus: '东校区' or '西校区'
        date_str: '2026.05.08'

    Returns:
        '东校区2026年5月8日教师值班记录.docx'
    """
    parts = date_str.split('.')
    year = int(parts[0])
    month = int(parts[1])
    day = int(parts[2])
    return f'{campus}{year}年{month}月{day}日教师值班记录.docx'

def find_document(base_dir, campus, date_str):
    """Find the correct Word document for given campus and date."""
    doc_name = build_doc_name(campus, date_str)
    doc_path = Path(base_dir) / campus / doc_name
    if doc_path.exists():
        return doc_path
    return None

def insert_image_to_doc(doc_path, image_path, timeslot_match, output_path=None):
    """Insert an image into the Word document at the matching timeslot row.

    Args:
        doc_path: Path to the .docx file
        image_path: Path to the image file to insert
        timeslot_match: dict from match_timeslot() with section/label
        output_path: Output path (defaults to overwrite doc_path)
    """
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(str(doc_path))
    section_name = timeslot_match['section']
    label = timeslot_match.get('label', '')

    # Find the matching row in tables.
    # Check each cell individually to avoid substring false matches
    # (e.g. "门口值日1" matching concatenated "门口值日10" from another cell).
    target_table = None
    target_row_idx = None

    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            row_text = ''
            cell_match = False
            for cell in row.cells:
                ct = cell.text.strip()
                row_text += ct + ' '
                if ct == section_name:
                    cell_match = True
            if not cell_match:
                # Fallback: section name embedded in cell (e.g. "课间值日(上午)")
                for cell in row.cells:
                    if section_name in cell.text:
                        cell_match = True
                        break
            if cell_match:
                if label and label not in row_text:
                    continue
                target_table = table
                target_row_idx = row_idx
                break
        if target_table:
            break

    if target_table is None:
        raise ValueError(f'Could not find row for section: {section_name}')

    # Insert image into the 3rd cell (情况记录 column)
    row = target_table.rows[target_row_idx]
    cell = row.cells[2]

    # Clear default empty paragraph
    cell.paragraphs[0].clear()

    # Add image (centered, 2.5 inches wide)
    run = cell.paragraphs[0].add_run()
    run.add_picture(str(image_path), width=Inches(2.5))

    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add timestamp subtitle
    timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
    p = cell.add_paragraph(f'[归档 {timestamp}]')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    save_path = Path(output_path) if output_path else Path(doc_path)
    doc.save(str(save_path))
    return save_path

# === CLI Entry Point ===
if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='Match time and insert image into duty log document')
    parser.add_argument('--image', required=True, help='Path to image file')
    parser.add_argument('--hhmm', required=True, help='Time in HH:MM format')
    parser.add_argument('--date', required=True, help='Date in YYYY.MM.DD format')
    parser.add_argument('--campus', required=True, help='东校区 or 西校区')
    parser.add_argument('--base-dir', required=True, help='Root directory of duty log documents')
    parser.add_argument('--dry-run', action='store_true', help='Print match without inserting')
    args = parser.parse_args()

    # 1. Match timeslot
    match = match_timeslot(args.hhmm)
    if match is None:
        print(f'ERROR: Could not match time {args.hhmm} to any timeslot')
        exit(1)

    print(f'Matched: {match["section"]} ({match.get("label", "N/A")})')

    # 2. Find document
    doc_path = find_document(args.base_dir, args.campus, args.date)
    if doc_path is None:
        doc_name = build_doc_name(args.campus, args.date)
        print(f'ERROR: Document not found: {args.base_dir}/{args.campus}/{doc_name}')
        exit(2)

    print(f'Document: {doc_path}')

    if args.dry_run:
        print('Dry run -- no changes made')
        exit(0)

    # 3. Insert image
    result = insert_image_to_doc(doc_path, args.image, match)
    print(f'Image inserted. Saved to: {result}')
